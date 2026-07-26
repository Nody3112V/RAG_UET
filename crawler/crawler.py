import os
import json
import requests
import time
import random
import pdfplumber
import pytesseract
import re
import hashlib
import unicodedata
from PIL import ImageOps, Image
import cv2
import numpy as np
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse, urldefrag, unquote, quote
from docx import Document
from pdf2image import convert_from_path
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# Tạo thư mục tạm để lưu file PDF, DOCX tải về
LUU_TAM = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "data", "raw", "file_tam")
# Kiểm tra thư mục tạm tồn tại chưa?
if not os.path.exists(LUU_TAM):
    os.makedirs(LUU_TAM)

URL = "https://handbook.uet.vnu.edu.vn/"
MIEN = "handbook.uet.vnu.edu.vn"

# Tạo một session requests với retry để xử lý các lỗi tạm thời khi tải trang web hoặc file
session = requests.Session()
session.headers.update({
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                  "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
})
retry = Retry(total=3, backoff_factor=1, status_forcelist=[429, 500, 502, 503, 504])
session.mount("https://", HTTPAdapter(max_retries=retry))
session.mount("http://", HTTPAdapter(max_retries=retry))

# chuẩn hoá các URL giống nhau nhưng khác nhau về định dạng, ví dụ: https://example.com và https://example.com/ sẽ được chuẩn hoá thành https://example.com
def normalize_url(url):
    parsed = urlparse(url)
    path = unicodedata.normalize('NFC', parsed.path)
    clean_url = f"{parsed.scheme}://{parsed.netloc}{path}"
    # Loại bỏ dấu gạch chéo ở cuối URL
    return clean_url.rstrip('/')

# Chuyển nội dung pdf scan sang văn bản sử dụng OCR
def chuyen_pdf_scan_sang_van_ban(duong_dan_file):
    text_ocr = ""
    try:
        print(f"Đang chuyển PDF scan sang văn bản: {duong_dan_file}")
        # Sử dụng thư viện pdf2image để chuyển PDF scan sang hình ảnh
        images = convert_from_path(duong_dan_file, dpi = 300)

        # Duyệt qua từng hình ảnh
        for stt, image in enumerate(images):
            text = pytesseract.image_to_string(image, lang='vie')  # sử dụng ngôn ngữ tiếng Việt
            if text:
                text_ocr += text + "\n"
                print(f"Thu thập thành công trang {stt + 1} của tài liệu PDF bằng OCR")
        return text_ocr

    except Exception as e:
        print(f"Lỗi khi chuyển PDF scan sang hình ảnh: {e}")
        return None

# Kiểm tra xem văn bản có phải là rác hay không, dựa trên tỷ lệ ký tự chữ và số trên tổng chiều dài văn bản
def is_garbage_text(text):
    if not text.strip(): return True
    # Tính tỷ lệ ký tự chữ và số trên tổng chiều dài
    alphanumeric_count = sum(c.isalnum() for c in text)
    if len(text) > 0 and (alphanumeric_count / len(text)) < 0.4:
        return True # Dưới 40% là chữ hoặc số thì coi là rác
    return False

# Tiền xử lý ảnh để tăng độ chính xác của OCR, bao gồm chuyển sang grayscale và áp dụng adaptive thresholding
def tien_xu_ly_anh(image):
    try:
        img = np.array(image.convert("L"))  # grayscale
        img = cv2.adaptiveThreshold(img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 15)
        return Image.fromarray(img)
    except Exception as e:
        print(f"Lỗi tiền xử lý ảnh: {e}")
        return image

# Thay thế các ký tự đặc biệt, khoảng trắng thừa và dòng trống thừa trong văn bản
def lam_sach_van_ban(text):
    text = text.replace('\f', '\n')
    text = re.sub(r'\n{3,}', '\n\n', text)          # gộp nhiều dòng trống
    text = re.sub(r'[ \t]{2,}', ' ', text)           # gộp khoảng trắng thừa
    return text.strip()

# tải file từ một URL, trích xuất nội dung văn bản từ file PDF hoặc DOCX, sau đó xóa file tạm và trả về nội dung văn bản đã trích xuất
# input: đường dẫn URL của file cần trích xuất nội dung
# output: nội dung văn bản đã trích xuất từ file PDF hoặc DOCX, nếu có lỗi thì trả về None
def truy_xuat_van_ban_tu_file_url(file_url):
    try:
        # Lấy dữ liệu file từ url
        du_lieu = session.get(file_url, stream=True, timeout = 10)
        # Kiểm tra xem lấy dữ liệu thành công chưa
        if du_lieu.status_code != 200:
            print(f"Lỗi tải file: {file_url}, mã lỗi: {du_lieu.status_code}")
            return None, None
            
        # lấy tên file từ url
        ten_file = unquote(os.path.basename(urlparse(file_url).path))
        if not ten_file:
            ten_file = "file_dinh_kem"
            
        content_type = du_lieu.headers.get('Content-Type', '').lower()
        content_disposition = du_lieu.headers.get('Content-Disposition', '')
        
        if 'filename=' in content_disposition:
            matches = re.findall(r'filename=["\']?([^"\';]+)["\']?', content_disposition)
            if matches:
                ten_file = matches[0]
                
        # Bổ sung đuôi file nếu thiếu, dựa vào Content-Type
        if not (ten_file.lower().endswith('.pdf') or ten_file.lower().endswith('.docx') or ten_file.lower().endswith('.doc')):
            if 'application/pdf' in content_type:
                ten_file += '.pdf'
            elif 'wordprocessingml.document' in content_type or 'msword' in content_type:
                ten_file += '.docx'

        # tạo đường dẫn đầy đủ để lưu vào file_tam bằng hash của URL để tránh trùng lặp
        hash_name = hashlib.md5(file_url.encode('utf-8')).hexdigest()
        ext = os.path.splitext(ten_file)[1]
        ten_file_tam = f"{hash_name}{ext}"
        duong_dan_file = os.path.join(LUU_TAM, ten_file_tam)

        # tạo (mở) file để lưu dữ liệu
        with open(duong_dan_file, "wb") as f:
            # ghi từng chunk dữ liệu (1024 byte) vào file
            for chunk in du_lieu.iter_content(chunk_size=1024):
                f.write(chunk)
        
        text_trich_xuat = ""

        # PDF
        if ten_file.lower().endswith(".pdf"):
            try:
                # Convert PDF sang ảnh một lần duy nhất ở đầu
                try:
                    images = convert_from_path(duong_dan_file, dpi=300)
                except Exception:
                    images = []

                # Mở file PDF và trích xuất văn bản từ từng trang
                with pdfplumber.open(duong_dan_file) as pdf:
                    for stt, trang in enumerate(pdf.pages):
                        text_trang = trang.extract_text()

                        # Trích xuất bảng biểu bằng pdfplumber và chuyển thành Markdown
                        tables = trang.extract_tables()
                        md_tables = ""
                        if tables:
                            for table in tables:
                                if not table: continue
                                md_tables += "\n"
                                for row_idx, row in enumerate(table):
                                    cleaned_row = [str(cell).replace('\n', ' ').replace('|', '-').strip() if cell else "" for cell in row]
                                    md_tables += "| " + " | ".join(cleaned_row) + " |\n"
                                    if row_idx == 0:
                                        md_tables += "|" + "|".join(["---"] * len(cleaned_row)) + "|\n"
                                md_tables += "\n"

                        # Nếu trang chỉ có chữ VÀ không phải là rác
                        if text_trang and text_trang.strip() and not is_garbage_text(text_trang):
                            text_trich_xuat += text_trang + "\n"
                            if md_tables:
                                text_trich_xuat += md_tables
                        elif stt < len(images):
                            # Nếu trang không có chữ hoặc là rác -> OCR 
                            print(f"Trang {stt+1} của file {ten_file} không có chữ hoặc là rác -> chạy OCR")
                            img_processed = tien_xu_ly_anh(images[stt])
                            text_ocr = pytesseract.image_to_string(img_processed, lang='vie', config='--oem 1 --psm 6') 
                            text_trich_xuat += text_ocr + "\n"
                            if md_tables:
                                text_trich_xuat += md_tables
                            print(f"Thu thập thành công trang {stt+1} của file {ten_file} bằng OCR")
            except Exception as e:
                print(f"Lỗi khi trích xuất PDF: {e}")
                return None, None
        
        # DOCX
        elif ten_file.lower().endswith(".docx"):
            try:
                # Mở file DOCX và trích xuất văn bản từ từng đoạn
                doc = Document(duong_dan_file)
                for doan in doc.paragraphs:
                    if doan.text.strip():
                        text_trich_xuat += doan.text.strip() + "\n"
                
                # Trích xuất văn bản từ bảng biểu (tables) và chuyển thành Markdown
                for table in doc.tables:
                    text_trich_xuat += "\n"
                    for row_idx, row in enumerate(table.rows):
                        row_data = []
                        for cell in row.cells:
                            text_cell = cell.text.strip().replace("\n", " ").replace("|", "-")
                            row_data.append(text_cell)
                        text_trich_xuat += "| " + " | ".join(row_data) + " |\n"
                        if row_idx == 0:
                            text_trich_xuat += "|" + "|".join(["---"] * len(row_data)) + "|\n"
                    text_trich_xuat += "\n"
            except Exception as e:
                print(f"Lỗi khi trích xuất DOCX: {e}")
                return None, None
        
        # Nếu định dạng file khác PDF hoặc DOCX, thông báo không hỗ trợ
        else:
            if ten_file.lower().endswith(".doc"):
                print(f"WARNING: Bỏ qua file .doc cũ (cần chuyển sang .docx): {ten_file}")
                with open(os.path.join(LUU_TAM, "loi_crawl.log"), "a", encoding="utf-8") as f_log:
                    f_log.write(f"Bỏ qua file .doc: {file_url}\n")
            else:
                print(f"Định dạng file không được hỗ trợ: {ten_file}")
            if os.path.exists(duong_dan_file):
                os.remove(duong_dan_file)
            return None, None
        
        # Nếu nội dung trích xuất từ file rỗng, thử OCR nếu là PDF scan
        if not text_trich_xuat.strip():
            print(f"Nội dung trích xuất từ file {ten_file} rỗng, thử OCR nếu là PDF scan")
            if ten_file.lower().endswith(".pdf"):
                text_ocr = chuyen_pdf_scan_sang_van_ban(duong_dan_file)
                if text_ocr and text_ocr.strip():
                    text_trich_xuat = text_ocr
                else:
                    print(f"Không thể trích xuất nội dung từ file PDF scan: {ten_file}")
                    if os.path.exists(duong_dan_file):
                        os.remove(duong_dan_file)  # xóa file tạm
                    return None, None

        # Xóa file tạm sau khi trích xuất xong
        if os.path.exists(duong_dan_file):
            os.remove(duong_dan_file)

        text_trich_xuat = lam_sach_van_ban(text_trich_xuat)
        return text_trich_xuat, ten_file
    except Exception as e:
        print(f"Lỗi khi xử lý file {file_url}: {e}")
        return None, None
    
# Thu thập nội dung của một trang web, bao gồm tiêu đề, nội dung văn bản và nội dung từ các file PDF/DOCX đính kèm
# Input: URL của trang web cần thu thập
# Output: một dictionary chứa URL, tiêu đề, nội dung văn bản và nội dung từ các file đính kèm, nếu có lỗi thì trả về None
def trich_xuat_noi_dung_trang_web(url, cache_van_ban):
    try:
        # Gửi yêu cầu lấy nội dung trang web
        yeu_cau = session.get(url, timeout = 10)
        # Nếu không thành công (mã lỗi khác 200), thông báo lỗi và trả về None
        if yeu_cau.status_code != 200:
            print(f"Lỗi khi truy cập trang web {url}, mã lỗi: {yeu_cau.status_code}")
            return None 

        # Chuyển yêu cầu thành đối tượng BeautifulSoup để phân tích HTML
        soup = BeautifulSoup(yeu_cau.content, "html.parser")

        # Lấy nội dung tiêu đề của web
        title_tag = soup.find('title')
        if title_tag:
            title = title_tag.text.strip()
        else:
            title = "Không có tiêu đề"

        # Tìm, lấy nội dung chính của trang web 
        lay_noi_dung = soup.find('article') or soup.find('main') or soup.find('div', class_='entry-content') or soup.find('body')
        # nối tất cả đoạn văn từ các thẻ cách nhau bằng dấu xuống dòng
        if lay_noi_dung:
            # xóa tất cả các thẻ script và style để tránh lấy nội dung không mong muốn
            for temp in lay_noi_dung.find_all(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                temp.decompose()

            # Danh sách các từ khóa nhận diện class rác
            tu_khoa_rac = [
                'menu', 'footer', 'sidebar', 'widget', 'nav', 'breadcrumb', 'header',
                'related-posts', 'share-buttons', 'comment', 'pagination', 'cookie-notice'
            ]

            # tạo một biểu thức chính quy để tìm các từ khóa rác trong class
            pattern_rac = re.compile('|'.join(tu_khoa_rac), re.IGNORECASE)

            # Tìm và loại bỏ các thẻ có class chứa từ khóa rác
            for rac in lay_noi_dung.find_all(class_=pattern_rac):
                rac.decompose()

            noi_dung = lay_noi_dung.get_text(separator="\n", strip=True)
        else:
            noi_dung = ""

        # trích xuất các liên kết đến file PDF hoặc DOCX từ trang web
        van_ban_dinh_kem = []
        link_thuong = []  # tập hợp các liên kết trong trang web

        # a là các thẻ chứa liên kết, href là thuộc tính chứa đường dẫn của liên kết
        links = soup.find_all('a', href=True)
        for link in links:

            # lấy đường dẫn đầy đủ của liên kết
            href = link['href']

            # lọc bỏ các liên kết không hợp lệ, bao gồm:
            # liên kết mailto: hoặc tel: hoặc javascript:
            href_lower = href.lower()
            if href_lower.startswith(('mailto:', 'tel:', 'javascript:')) or "@" in href_lower or not href.strip():
                continue

            # Mã hoá URL để tránh lỗi khi có ký tự đặc biệt
            href = quote(href, safe=':/?#[]@!$&\'()*+,;=%')

            full_url = urljoin(url, href)

            # kiểm tra xem liên kết có phải là file PDF hoặc DOCX không
            is_file = False
            full_url_lower = full_url.lower()
            if full_url_lower.endswith(('.pdf', '.docx', '.doc')):
                is_file = True
            else:
                # Kiểm tra fallback bằng HEAD request nếu URL có dấu hiệu là file tải về
                tu_khoa_file = ['/download', '/upload', 'wp-content/uploads', 'file=', 'id=', 'attachment']
                if any(tk in full_url_lower for tk in tu_khoa_file):
                    try:
                        head_req = session.head(full_url, timeout=5, allow_redirects=True)
                        ct = head_req.headers.get('Content-Type', '').lower()
                        if 'application/pdf' in ct or 'word' in ct or 'officedocument' in ct:
                            is_file = True
                    except Exception:
                        pass

            if is_file:
                if full_url in cache_van_ban:
                    item = cache_van_ban[full_url]
                    if item:
                        van_ban_dinh_kem.append(item)
                else:
                    print("phát hiện file đính kèm:", full_url)
                    # trích xuất nội dung văn bản từ file đính kèm
                    van_ban, ten_file_goc = truy_xuat_van_ban_tu_file_url(full_url)
                    # Đảm bảo rằng văn bản trích xuất không phải là None trước khi nối vào van_ban_dinh_kem
                    if van_ban:
                        item = {
                            "loai": "file_dinh_kem",
                            "ten_file": ten_file_goc,
                            "url_file": full_url,
                            "noi_dung": van_ban
                        }
                        van_ban_dinh_kem.append(item)
                        cache_van_ban[full_url] = item
                    else:
                        cache_van_ban[full_url] = None  # đánh dấu là đã thử tải nhưng lỗi
            else:
                link_thuong.append(full_url)
        return {
            "loai": "trang_web",
            "url": url,
            "title": title,
            "noi_dung": noi_dung,
            "van_ban_dinh_kem": van_ban_dinh_kem,
            "link_thuong": link_thuong
        }
    except Exception as e:
        print(f"Lỗi khi truy cập trang web {url}: {e}")
        with open(os.path.join(LUU_TAM, "loi_crawl.log"), "a", encoding="utf-8") as f_log:
            f_log.write(f"Lỗi trang web: {url} - {str(e)}\n")
        return None

# Bắt đầu từ một URL gốc, thu thập nội dung của trang hiện tại, sau đó tìm các liên kết nội bộ trên trang để tiếp tục thu thập nội dung của các trang con, lặp lại quá trình này cho đến khi không còn trang mới hoặc đạt số lượng trang tối đa cần crawl.
def thu_thap(start_url, max_pages = None):
    checkpoint_file = os.path.join(LUU_TAM, "checkpoint_crawl.json")
    
    if os.path.exists(checkpoint_file):
        print("Tìm thấy file checkpoint, đang tải lại tiến độ...")
        with open(checkpoint_file, "r", encoding="utf-8") as f:
            data = json.load(f)
            visiteds_urls = set(data.get("visiteds_urls", []))
            cache_van_ban = data.get("cache_van_ban", {})
            queue = data.get("queue", [])
            results = data.get("results", [])
    else:
        visiteds_urls = set()  # tập hợp các URL đã truy cập
        cache_van_ban = {}  # dict cache nội dung các file PDF/DOCX đã tải
        start_url = normalize_url(start_url)  # chuẩn hoá URL gốc
        queue = [start_url]  # hàng đợi các URL cần truy cập
        queue_set = set(queue) # Set để kiểm tra O(1)
        results = []  # danh sách kết quả thu thập
        
        # Đọc sitemap.xml
        print("Đang quét cấu trúc Sitemap...")
        sitemap_url = urljoin(start_url, "sitemap.xml")
        try:
            sitemap_req = session.get(sitemap_url, timeout=10)
            if sitemap_req.status_code == 200:
                sitemap_soup = BeautifulSoup(sitemap_req.content, "xml")
                for loc in sitemap_soup.find_all("loc"):
                    loc_url = normalize_url(loc.text)
                    if urlparse(loc_url).netloc == MIEN and loc_url not in queue_set:
                        queue.append(loc_url)
                        queue_set.add(loc_url)
                print(f"Đã tìm thấy {len(queue)} URL từ sitemap.")
        except Exception as e:
            print(f"Không lấy được sitemap: {e}")

    # duyệt từng URL bằng BFS
    while queue:
        # Nếu có giới hạn số lượng trang cần crawl và đã đạt giới hạn thì dừng lại
        if max_pages is not None and len(visiteds_urls) >= max_pages:
            print(f"Đã đạt giới hạn số lượng trang cần crawl: {max_pages}")
            break

        current_url = queue.pop(0)

        if current_url in visiteds_urls:
            continue

        print(f"Đang truy cập: {current_url}")
        visiteds_urls.add(current_url)

        time.sleep(random.uniform(1, 3))  # tạm dừng từ 1 đến 3 giây để tránh bị chặn

        page_data = trich_xuat_noi_dung_trang_web(current_url, cache_van_ban)
        # kiểm tra nếu page_data không rỗng và có nội dung hoặc văn bản đính kèm, có thì thêm vào kết quả
        if page_data:
            next_links = page_data.pop("link_thuong", [])

            if page_data.get("noi_dung") or page_data.get("van_ban_dinh_kem"):
                results.append(page_data)
                
                # In thông báo thu thập thành công nổi bật
                title = page_data.get('title', 'Không có tiêu đề')
                num_files = len(page_data.get("van_ban_dinh_kem") or [])
                print(f"\n✅ [THÀNH CÔNG] Đã thu thập: {title}")
                print(f"   🔗 URL: {current_url}")
                if num_files > 0:
                    print(f"   📎 Kèm theo {num_files} file đính kèm (PDF/DOCX).")

            # đưa các liên kết nội bộ chưa truy cập vào hàng đợi để tiếp tục thu thập 
            for next_link in next_links:
                # chuẩn hoá liên kết trước khi thêm vào hàng đợi
                next_link = normalize_url(next_link)
                # chỉ thêm các liên kết nội bộ (cùng miền) vào hàng đợi
                if urlparse(next_link).netloc == MIEN and next_link not in queue_set and next_link not in visiteds_urls:
                    queue.append(next_link)
                    queue_set.add(next_link)
                    
        # Lưu checkpoint mỗi 20 trang
        if len(visiteds_urls) > 0 and len(visiteds_urls) % 20 == 0:
            print(f"Đang lưu checkpoint tại trang thứ {len(visiteds_urls)}...")
            try:
                with open(checkpoint_file, "w", encoding="utf-8") as f:
                    json.dump({
                        "visiteds_urls": list(visiteds_urls),
                        "cache_van_ban": cache_van_ban,
                        "queue": queue,
                        "results": results
                    }, f, ensure_ascii=False, indent=4)
            except Exception as e:
                print(f"Lỗi khi lưu checkpoint: {e}")

    # Thu thập xong toàn bộ, xóa file checkpoint
    if os.path.exists(checkpoint_file):
        try:
            os.remove(checkpoint_file)
        except Exception:
            pass

    return results

if __name__ == "__main__":
    final_data = thu_thap(URL)
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "data", "raw", "ket_qua_crawl.json")
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(final_data, f, ensure_ascii=False, indent=4)
    print(f"Đã lưu kết quả vào file {output_path}")